from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.common.keys import Keys
from docx import Document
from docx.shared import Inches
import time
import os

# List of URLs and matching search terms (1:1 mapping)
urls_to_test = [
    "https://softwaredevelopmentsolution.com/Access/RoleTypeSetup",
    "https://softwaredevelopmentsolution.com/Access/UserSetup",
    "https://softwaredevelopmentsolution.com/DocumentSettings/AccountingDocumentPrefix",
    "https://softwaredevelopmentsolution.com/Accounting/PaymentTerm",
    "https://softwaredevelopmentsolution.com/Accounting/AccountType",
    "https://softwaredevelopmentsolution.com/Accounting/AccountDetailType",
    "https://softwaredevelopmentsolution.com/Accounting/Group",
    "https://softwaredevelopmentsolution.com/Accounting/Account",
    "https://softwaredevelopmentsolution.com/Accounting/PaymentSurcharge",
    "https://softwaredevelopmentsolution.com/Accounting/MiscellaneousItem",
    "https://softwaredevelopmentsolution.com/Accounting/TaxType",
    "https://softwaredevelopmentsolution.com/Accounting/JournalEntry",
    "https://softwaredevelopmentsolution.com/Accounting/CashEntry",
    "https://softwaredevelopmentsolution.com/Accounting/BankEntry",
    "https://softwaredevelopmentsolution.com/Accounting/PurchaseEntry",
    "https://softwaredevelopmentsolution.com/Accounting/OpeningEntry",
    "https://softwaredevelopmentsolution.com/Accounting/PaymentMatching",
    "https://softwaredevelopmentsolution.com/Accounting/PaymentRequest",
    "https://softwaredevelopmentsolution.com/Accounting/PaymentReceived",
    "https://softwaredevelopmentsolution.com/Accounting/BankDeposit",
    "https://softwaredevelopmentsolution.com/Accounting/Reconcile",
    "https://softwaredevelopmentsolution.com/Contacts/CustVendContactTypes",
    "https://softwaredevelopmentsolution.com/Contacts/CustomerVendor",
    "https://softwaredevelopmentsolution.com/DocumentSettings/InventoryDocumentPrefix",
    "https://softwaredevelopmentsolution.com/Inventory/Characteristics",
    "https://softwaredevelopmentsolution.com/Inventory/CharacteristicUnits",
    "https://softwaredevelopmentsolution.com/Inventory/CharacteristicValues",
    "https://softwaredevelopmentsolution.com/Inventory/Categories",
    "https://softwaredevelopmentsolution.com/Inventory/ProductUnit",
    "https://softwaredevelopmentsolution.com/Inventory/ProductDiscountLevel",
    "https://softwaredevelopmentsolution.com/Inventory/Products",
    "https://softwaredevelopmentsolution.com/Inventory/ServiceCategory",
    "https://softwaredevelopmentsolution.com/Inventory/ServiceUnit",
    "https://softwaredevelopmentsolution.com/Inventory/Service",
    "https://softwaredevelopmentsolution.com/Inventory/StockTransfer",
    "https://softwaredevelopmentsolution.com/Inventory/StockAdjustment",
    "https://softwaredevelopmentsolution.com/Inventory/CycleCount",
    "https://softwaredevelopmentsolution.com/DocumentSettings/PurchaseDocumentPrefix",
    "https://softwaredevelopmentsolution.com/Purchase/Vendors",
    #"https://softwaredevelopmentsolution.com/Purchase/PurchaseOrder",
    "https://softwaredevelopmentsolution.com/Purchase/Receive",
    "https://softwaredevelopmentsolution.com/Purchase/GoodsReturn",
    "https://softwaredevelopmentsolution.com/Purchase/Invoice",
    "https://softwaredevelopmentsolution.com/Purchase/PurchaseReturn",
    "https://softwaredevelopmentsolution.com/DocumentSettings/ManufacturingDocumentPrefix",
    "https://softwaredevelopmentsolution.com/Manufacturing/WorkstationType",
    "https://softwaredevelopmentsolution.com/Manufacturing/Workstation",
    "https://softwaredevelopmentsolution.com/Manufacturing/Operation",
    "https://softwaredevelopmentsolution.com/Manufacturing/Process",
    "https://softwaredevelopmentsolution.com/Manufacturing/Routing",
    "https://softwaredevelopmentsolution.com/Manufacturing/Bom",
    "https://softwaredevelopmentsolution.com/Manufacturing/Production",
    "https://softwaredevelopmentsolution.com/Manufacturing/JobCard",
    "https://softwaredevelopmentsolution.com/DocumentSettings/SalesDocumentPrefix",
    "https://softwaredevelopmentsolution.com/Sale/Customers",
    "https://softwaredevelopmentsolution.com/Sale/SaleQuote",
    "https://softwaredevelopmentsolution.com/Sale/SaleOrder",
    "https://softwaredevelopmentsolution.com/Sale/SaleDeliveryReturn",
    "https://softwaredevelopmentsolution.com/Sale/SaleInvoice",
    "https://softwaredevelopmentsolution.com/Sale/SaleReturn",
    "https://softwaredevelopmentsolution.com/Shipping/ShippingCategories",
    "https://softwaredevelopmentsolution.com/Shipping/ShippingTerms",
    "https://softwaredevelopmentsolution.com/Shipping/ShippingCarrier",
    "https://softwaredevelopmentsolution.com/Shipping/ShipmentPackage",
    "https://softwaredevelopmentsolution.com/Shipping/Delivery",
    "https://softwaredevelopmentsolution.com/Shipping/Picking",
    "https://softwaredevelopmentsolution.com/Shipping/Shipment",
    "https://softwaredevelopmentsolution.com/DocumentSettings/HRDocumentPrefix",
    "https://softwaredevelopmentsolution.com/HR/Employees",
    "https://softwaredevelopmentsolution.com/HR/EmploymentTypes",
    "https://softwaredevelopmentsolution.com/HR/Designation",
    "https://softwaredevelopmentsolution.com/HR/SalaryComponents",
    "https://softwaredevelopmentsolution.com/HR/SalaryStructures",
    "https://softwaredevelopmentsolution.com/HR/PayrollSchedule",
    "https://softwaredevelopmentsolution.com/HR/Activity",
    "https://softwaredevelopmentsolution.com/HR/ActivityType",
    "https://softwaredevelopmentsolution.com/HR/Project",
    "https://softwaredevelopmentsolution.com/HR/Job",
    "https://softwaredevelopmentsolution.com/HR/Holidays",
    "https://softwaredevelopmentsolution.com/HR/HolidayPolicy",
    "https://softwaredevelopmentsolution.com/HR/Timesheet",
    "https://softwaredevelopmentsolution.com/Settings/SalesTax",
    "https://softwaredevelopmentsolution.com/Settings/ScreenTerms",
    "https://softwaredevelopmentsolution.com/Settings/LocationType",
    "https://softwaredevelopmentsolution.com/Settings/Location",
    "https://softwaredevelopmentsolution.com/Settings/LocationBin"

]

search_terms = [
    "Finance and Accounting",                     
    "Christopher Smith",                    
    "PIYUSH",
    "Cash Payment",
    "Admin Expenses",
    "Accounts Payables",
    "Financial Groups",
    "Accounts Receivable",
    "Let Payment Surcharges",
    "Office supplies to accessories",
    "Capital Gain Tax",
    "1,200,000",
    "40,000",
    "10,500",
    "06/17/2024",
    "10/01/2024",
    "BroToMotive",
    "Jems OLA Customer",
    "32,619",
    "PDFPDFERDT000015",
    "Bank of America",
    "Purchase Manager",
    "Customer Mails",
    "Stock Transfer",
    "Comfort",
    "Battery Capacity",
    "Technology Integration",
    "City/Urban",
    "Container",
    "Promotional Discounts",
    "Bike Tire 100",
    "Customization",
    "Customization Service",
    "Component Upgrades and Custom",
    "09/05/2024",
    "OLA Sell Shop",
    "CCC000026",
    "PR",
    "Jonny Deep",
    #"",
    "Theodore Adams",
    "08/13/2024",
    "Jems OLA Vendor",
    "OLA Electric Showroom (TW)",
    "Manufacturing Order",
    "Flexible Manufacturing System",
    "Flexible Manufacturing Systems of Tier",
    "Mango",
    "Manufacturing of Component",
    "Mango",
    "02/12/2025",
    "OLS PROD STOCK TEST",
    "Onboard charger",
    "SDN",
    "Regular Customer",
    "Bhopendra Customer",
    "Bhopendra Customer",
    "Bhopendra Customer",
    "Ola Shop001",
    "OLA S&P",
    "Standard Shipping",
    "By Truck",
    "Courier Services",
    "Cantener Of 10 Crate",
    "OLA MFR Vendor",
    "SON000083",
    "OLA S&P",
    "TS",
    "Henry Brown",
    "Part Time",
    "Manager",
    "Base Salary",
    "Fortnightly",
    "OTP PS",
    "Electric Maintenance",
    "Planning",
    "Piyush P Suthar",
    "Paint Job",
    "Halloween",
    "Test",
    "TEST EMP Smith",
    "69857",
    "Sales Return",
    "Point Of Sales",
    "OLA Electric",
    "EV Storage"

]

# Login credentials
email = "ola123@yopmail.com"
password = "ola@1234"

# Create the Word document
doc = Document()
doc.add_heading("Multi-Page, Single-Term Search Automation Report", level=1)

# Setup Edge WebDriver
options = webdriver.EdgeOptions()
options.use_chromium = True

driver = webdriver.Edge(options=options)
wait = WebDriverWait(driver, 10)

# Login function
def login():
    driver.get(urls_to_test[0])
    try:
        username_field = wait.until(EC.presence_of_element_located((By.ID, "Email")))
        password_field = wait.until(EC.presence_of_element_located((By.ID, "Password")))
        login_button = wait.until(EC.element_to_be_clickable((By.ID, "LoginSubmit")))

        username_field.send_keys(email)
        password_field.send_keys(password)
        login_button.click()

        wait.until(EC.presence_of_element_located((By.ID, "mobile-detailed-view")))
        print("✅ Login successful!")
    except Exception as e:
        print(f"⚠️ Login not required or already authenticated: {e}")

# Perform login once
login()

# Loop through each URL and matching search term
for index, (url, term) in enumerate(zip(urls_to_test, search_terms), start=1):
    try:
        driver.get(url)

        # Wait for the search input field
        search_input = wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, "input[type='search'].form-control.form-control-sm")))
        driver.execute_script("arguments[0].focus();", search_input)
        search_input.clear()
        search_input.send_keys(term)
        search_input.send_keys(Keys.RETURN)

        # Wait for search results to update
        time.sleep(2)

        # Screenshot name
        timestamp = time.strftime("%Y-%m-%d_%H-%M-%S")
        screenshot_name = f"screenshot_{index}_{timestamp}.png"
        driver.save_screenshot(screenshot_name)
        print(f"📸 Screenshot taken for URL: {url} | Term: {term}")

        # Add to Word document
        doc.add_heading(f"Page {index} - Search '{term}'", level=2)
        doc.add_paragraph(f"URL: {url}")
        doc.add_paragraph(f"Searched at: {timestamp}")
        doc.add_picture(screenshot_name, width=Inches(5))

    except Exception as e:
        print(f"❌ Error on {url} with term '{term}': {e}")
        doc.add_heading(f"Error on {url}", level=2)
        doc.add_paragraph(f"Search term: {term}")
        doc.add_paragraph(f"Error: {str(e)}")

# Save Word document
report_file = "Multi_Page_Single_Term_Search_Report.docx"
doc.save(report_file)
print(f"✅ Final report saved as: {report_file}")

# Close browser
driver.quit()
