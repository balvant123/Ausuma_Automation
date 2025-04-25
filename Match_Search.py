import tkinter as tk
from tkinter import ttk
import threading

# Example URL list
urls = {
    "Inventory - Products": "https://softwaredevelopmentsolution.com/Inventory/Products",
    "Accounting - Journal Entry": "https://softwaredevelopmentsolution.com/Accounting/JournalEntry",
    "Sales - Customers": "https://softwaredevelopmentsolution.com/Sale/Customers"
}

def run_automation(selected_url, search_term):
    from selenium import webdriver
    from selenium.webdriver.common.by import By
    from selenium.webdriver.support.ui import WebDriverWait
    from selenium.webdriver.support import expected_conditions as EC
    from selenium.webdriver.common.keys import Keys
    from docx import Document
    from docx.shared import Inches
    import time

    doc = Document()
    doc.add_heading("Search Automation Report", level=1)

    email = "ola123@yopmail.com"
    password = "ola@1234"

    options = webdriver.EdgeOptions()
    options.use_chromium = True
    driver = webdriver.Edge(options=options)
    wait = WebDriverWait(driver, 10)

    def login():
        driver.get(selected_url)
        try:
            username_field = wait.until(EC.presence_of_element_located((By.ID, "Email")))
            password_field = wait.until(EC.presence_of_element_located((By.ID, "Password")))
            login_button = wait.until(EC.element_to_be_clickable((By.ID, "LoginSubmit")))

            username_field.send_keys(email)
            password_field.send_keys(password)
            login_button.click()

            wait.until(EC.presence_of_element_located((By.ID, "mobile-detailed-view")))
            print("✅ Login successful!")
        except Exception as e:
            print(f"⚠️ Login not required or already authenticated: {e}")

    login()

    driver.get(selected_url)
    try:
        search_input = wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, "input[type='search'].form-control.form-control-sm")))
        search_input.clear()
        search_input.send_keys(search_term)
        search_input.send_keys(Keys.RETURN)
        time.sleep(2)

        timestamp = time.strftime("%Y-%m-%d_%H-%M-%S")
        screenshot_name = f"screenshot_{timestamp}.png"
        driver.save_screenshot(screenshot_name)

        doc.add_heading(f"Search '{search_term}'", level=2)
        doc.add_paragraph(f"URL: {selected_url}")
        doc.add_paragraph(f"Searched at: {timestamp}")
        doc.add_picture(screenshot_name, width=Inches(5))
        print("📸 Screenshot saved.")

    except Exception as e:
        doc.add_heading("Error", level=2)
        doc.add_paragraph(str(e))

    doc.save("Selected_URL_Search_Report.docx")
    driver.quit()
    print("✅ Report generated!")

# GUI
def on_submit():
    search_term = entry.get()
    if search_all_var.get():
        for url_name, selected_url in urls.items():
            threading.Thread(target=run_automation, args=(selected_url, search_term)).start()
    else:
        selected_text = combo.get()
        selected_url = urls[selected_text]
        threading.Thread(target=run_automation, args=(selected_url, search_term)).start()

root = tk.Tk()
root.title("Search Automation")

tk.Label(root, text="Select URL:").pack()
combo = ttk.Combobox(root, values=list(urls.keys()))
combo.pack()
combo.current(0)

tk.Label(root, text="Enter Search Term:").pack()
entry = tk.Entry(root)
entry.pack()

# ✅ Checkbox for "Search all URLs"
search_all_var = tk.BooleanVar()
tk.Checkbutton(root, text="Search all URLs", variable=search_all_var).pack()

tk.Button(root, text="Run", command=on_submit).pack()

root.mainloop()
